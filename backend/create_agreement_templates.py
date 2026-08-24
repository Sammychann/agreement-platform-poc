"""
=============================================================================
EXACT WORD AGREEMENT TEMPLATE GENERATOR
Based on exact images from 'Refer these images' folder:
  1. Direct Agreement Template-Customer ownership.docx
  2. Direct Agreement Template-Innoject Pro.docx
  3. Indirect Agreement Template-Customer Ownership.docx
  4. Indirect Agreement Template-Innoject Pro.docx
=============================================================================
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement
from pathlib import Path

from config import TEMPLATES_DIR


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner margins for a table cell in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table):
    """Apply standard clean black border to table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def remove_table_borders(table):
    """Remove all borders from a table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def set_cell_width(cell, width_inches):
    """Explicitly set cell width."""
    cell.width = Inches(width_inches)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def format_para(p, text="", font_name="Times New Roman", font_size=11, bold=False, italic=False, underline=False, align=None, space_after=Pt(3), space_before=Pt(0), line_spacing=1.15):
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.underline = underline
        run.font.name = font_name
        run.font.size = Pt(font_size)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    p.paragraph_format.line_spacing = line_spacing
    return p


def add_p(doc, text="", font_name="Times New Roman", font_size=11, bold=False, italic=False, underline=False, align=None, space_after=Pt(4), space_before=Pt(0), line_spacing=1.15):
    p = doc.add_paragraph()
    return format_para(p, text, font_name, font_size, bold, italic, underline, align, space_after, space_before, line_spacing)


def add_run_to_para(p, text, font_name="Times New Roman", font_size=11, bold=False, italic=False, underline=False):
    """Add a run to an existing paragraph with specified formatting."""
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.name = font_name
    run.font.size = Pt(font_size)
    return run


# =============================================================================
# PAGE 1: APPENDIX B OPPORTUNITY EVALUATION WORKSHEET (ALL 4 TEMPLATES)
# =============================================================================
def build_page_1_appendix_b(doc, is_indirect=False):
    # Header: Proprietary (Top Left)
    p_prop = add_p(doc, "Proprietary", font_size=9, bold=False, italic=False, space_after=Pt(12))

    # Titles
    add_p(doc, "MAH-PROC-100-12 Product Related Technical Equipment and Devices", font_size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_p(doc, "APPENDIX B", font_size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_p(doc, "Opportunity Evaluation Worksheet", font_size=10, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_p(doc, "(Local currency)", font_size=9, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(10))

    # ---- Main Details Table (3 rows, 2 cols) ----
    table1 = doc.add_table(rows=3, cols=2)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table1)

    for row in table1.rows:
        row.cells[0].width = Inches(3.6)
        row.cells[1].width = Inches(3.6)
        for cell in row.cells:
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)

    # Row 0: Customer Name & Location
    c00 = table1.cell(0, 0)
    p00 = c00.paragraphs[0]
    if is_indirect:
        format_para(p00, "Customer Name: {Distributor Name}\n{{Customer Name}}", font_size=9.5)
    else:
        format_para(p00, "Customer Name: {Customer Name}", font_size=9.5)

    c01 = table1.cell(0, 1)
    p01 = c01.paragraphs[0]
    format_para(p01, "Location: {Location}", font_size=9.5)

    # Row 1: Purpose (merge across both columns to span full width)
    c10 = table1.cell(1, 0)
    c11 = table1.cell(1, 1)
    c10.merge(c11)
    p10 = c10.paragraphs[0]
    format_para(p10, "Purpose for Providing Equipment:    To promote usage of Hatchery Vaccination", font_size=9.5)

    # Row 2: Sales Value & Margins
    c20 = table1.cell(2, 0)
    p20 = c20.paragraphs[0]
    format_para(p20, "Sales Value of Initial Purchase Order:\nAs Per financial Evaluation", font_size=9.5)

    c21 = table1.cell(2, 1)
    p21 = c21.paragraphs[0]
    format_para(p21, "Estimated Local Gross Margin of Annual Sales Value*: As Per financial Evaluation\n\nAdditional finance support from local / regional stakeholders - ", font_size=9.5)

    # ---- Equipment Table Section ----
    # "Equipment and/or Devices to be provided:" header
    p_eq_header = add_p(doc, "Equipment and/or Devices to be provided:", font_size=9.5, space_before=Pt(6), space_after=Pt(2))

    # Equipment table with Description / Quantity / FMV columns
    eq_table = doc.add_table(rows=4, cols=3)
    eq_table.style = 'Table Grid'
    eq_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(eq_table)

    for row in eq_table.rows:
        set_cell_width(row.cells[0], 2.5)
        set_cell_width(row.cells[1], 1.5)
        set_cell_width(row.cells[2], 1.5)
        for cell in row.cells:
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

    # Header row
    format_para(eq_table.cell(0, 0).paragraphs[0], "Description", font_size=9, bold=True, underline=True)
    format_para(eq_table.cell(0, 1).paragraphs[0], "Quantity", font_size=9, bold=True, underline=True)
    format_para(eq_table.cell(0, 2).paragraphs[0], "FMV", font_size=9, bold=True)

    # Equipment rows (placeholders for dynamic replacement)
    format_para(eq_table.cell(1, 0).paragraphs[0], "{Equipment Name}", font_size=9)
    format_para(eq_table.cell(1, 1).paragraphs[0], "{quantity}", font_size=9)
    format_para(eq_table.cell(1, 2).paragraphs[0], "", font_size=9)

    format_para(eq_table.cell(2, 0).paragraphs[0], "{Equipment Name}", font_size=9)
    format_para(eq_table.cell(2, 1).paragraphs[0], "{quantity}", font_size=9)
    format_para(eq_table.cell(2, 2).paragraphs[0], "", font_size=9)

    # Empty row for additional items
    format_para(eq_table.cell(3, 0).paragraphs[0], "", font_size=9)
    format_para(eq_table.cell(3, 1).paragraphs[0], "", font_size=9)
    format_para(eq_table.cell(3, 2).paragraphs[0], "", font_size=9)

    # Marker for dynamic equipment replacement
    p_eq_marker = add_p(doc, "{{appendix_b_equipment_table}}", font_size=1, space_after=Pt(0), space_before=Pt(0))

    # Total FAIR MARKET VALUE & ROI table
    roi_table = doc.add_table(rows=1, cols=2)
    roi_table.style = 'Table Grid'
    roi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(roi_table)

    set_cell_width(roi_table.cell(0, 0), 3.6)
    set_cell_width(roi_table.cell(0, 1), 3.6)
    set_cell_margins(roi_table.cell(0, 0), top=60, bottom=60, left=100, right=100)
    set_cell_margins(roi_table.cell(0, 1), top=60, bottom=60, left=100, right=100)

    format_para(roi_table.cell(0, 0).paragraphs[0], "Total FAIR MARKET VALUE of Equipment and/or Devices:", font_size=9)
    format_para(roi_table.cell(0, 1).paragraphs[0], "Return on Investment:\n[Gross Margin of Annual Sales - FMV of Equipment and/or Devices]÷ [FMV of Equipment and Devices]: %", font_size=8.5)

    # ROI note below table
    add_p(doc, "*Similar to any investment MAH may make, we would expect a reasonable Return on Investment (ROI) related to this investment in Equipment and/or Devices.", font_size=8.5, italic=True, space_before=Pt(4), space_after=Pt(8))

    # ---- Approvers Table ----
    table_app = doc.add_table(rows=5, cols=3)
    table_app.style = 'Table Grid'
    table_app.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_app)

    for row in table_app.rows:
        set_cell_width(row.cells[0], 0.8)
        set_cell_width(row.cells[1], 3.0)
        set_cell_width(row.cells[2], 3.4)
        for cell in row.cells:
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

    # Header Row
    format_para(table_app.cell(0, 0).paragraphs[0], "", font_size=9)
    format_para(table_app.cell(0, 1).paragraphs[0], "", font_size=9)
    format_para(table_app.cell(0, 2).paragraphs[0], "If signed by hand, a signature date is required", font_size=8.5, bold=True, italic=True)

    # Approvers vertical cell (merge col 0 across rows 1-4)
    cell_app_vert = table_app.cell(1, 0)
    for r_idx in range(2, 5):
        cell_app_vert.merge(table_app.cell(r_idx, 0))
    p_v = cell_app_vert.paragraphs[0]
    format_para(p_v, "Approvers", font_size=9, bold=True)

    # Rows 1 to 4
    app_roles = [
        ("Initiator/Business Owner", "{Initiator Name and Date}"),
        ("Next level manager (with appropriate GOA)", "{Manager Name and Date}"),
        ("Regional Finance\n(for leased/ rental user agreement for a fee)", ""),
        ("Legal\n(when dominant Products are involved)", "")
    ]

    for idx, (role, sig) in enumerate(app_roles):
        row_num = idx + 1
        c_role = table_app.cell(row_num, 1)
        format_para(c_role.paragraphs[0], role, font_size=8.5, bold=True, italic=True)
        c_sig = table_app.cell(row_num, 2)
        format_para(c_sig.paragraphs[0], sig, font_size=8.5)

    doc.add_page_break()


# =============================================================================
# SIGNATURE SECTION — Matches reference images (side-by-side blocks, no table)
# =============================================================================
def build_signature_section(doc, left_party_name="{Customer Name}", right_party_name="Intervet India Private Limited"):
    """Build the signature section matching reference images.
    Uses a borderless 2-column table to create side-by-side layout.
    """
    add_p(doc, "IN WITNESS WHEREOF, the Parties have caused this Form to be duly executed as of the Effective Date.", font_size=10, space_after=Pt(16))

    # Use a borderless table for side-by-side layout
    sig_table = doc.add_table(rows=7, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(sig_table)

    for row in sig_table.rows:
        set_cell_width(row.cells[0], 3.5)
        set_cell_width(row.cells[1], 3.5)
        for cell in row.cells:
            set_cell_margins(cell, top=40, bottom=40, left=100, right=100)

    # Row 0: "Signed for and on behalf of:" headers
    format_para(sig_table.cell(0, 0).paragraphs[0], "Signed for and on behalf of:", font_size=10)
    format_para(sig_table.cell(0, 1).paragraphs[0], "Signed for and on behalf of:", font_size=10)

    # Row 1: Party names
    format_para(sig_table.cell(1, 0).paragraphs[0], left_party_name, font_size=10)
    format_para(sig_table.cell(1, 1).paragraphs[0], right_party_name, font_size=10, underline=True)

    # Row 2: Blank space for signatures + signature image markers
    p_cust_sig = sig_table.cell(2, 0).paragraphs[0]
    format_para(p_cust_sig, "\n\n{{customer_signature}}", font_size=9.5, space_after=Pt(2))

    p_intervet_sig = sig_table.cell(2, 1).paragraphs[0]
    format_para(p_intervet_sig, "\n\n{{intervet_signature}}", font_size=9.5, space_after=Pt(2))

    # Row 3: Signature lines (horizontal rule using underscores)
    format_para(sig_table.cell(3, 0).paragraphs[0], "________________________________", font_size=10, space_after=Pt(4))
    format_para(sig_table.cell(3, 1).paragraphs[0], "________________________________", font_size=10, space_after=Pt(4))

    # Row 4: Name fields
    p_name_left = sig_table.cell(4, 0).paragraphs[0]
    add_run_to_para(p_name_left, "Name: ", font_size=10)
    add_run_to_para(p_name_left, "{Receiver Name}", font_size=10)
    p_name_left.paragraph_format.space_after = Pt(2)

    p_name_right = sig_table.cell(4, 1).paragraphs[0]
    add_run_to_para(p_name_right, "Name: ", font_size=10)
    add_run_to_para(p_name_right, "{Name}", font_size=10)
    p_name_right.paragraph_format.space_after = Pt(2)

    # Row 5: Title fields
    p_title_left = sig_table.cell(5, 0).paragraphs[0]
    add_run_to_para(p_title_left, "Title: ", font_size=10)
    add_run_to_para(p_title_left, "{Title of receiver}", font_size=10)
    p_title_left.paragraph_format.space_after = Pt(2)

    p_title_right = sig_table.cell(5, 1).paragraphs[0]
    add_run_to_para(p_title_right, "Title: ", font_size=10)
    add_run_to_para(p_title_right, "{Title}", font_size=10)
    p_title_right.paragraph_format.space_after = Pt(2)

    # Row 6: Date fields
    format_para(sig_table.cell(6, 0).paragraphs[0], "Date:", font_size=10, space_after=Pt(2))
    format_para(sig_table.cell(6, 1).paragraphs[0], "Date:", font_size=10, space_after=Pt(2))


# =============================================================================
# EXHIBIT A & EXHIBIT X (COMMON TO TEMPLATES)
# =============================================================================
def build_exhibit_a(doc, title_subtitle="Products and Quantities"):
    doc.add_page_break()
    add_p(doc, "Confidential", font_size=9, space_after=Pt(18))
    add_p(doc, "EXHIBIT A", font_size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_p(doc, title_subtitle, font_size=10, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(16))

    # Exhibit A Dynamic Equipment Marker
    add_p(doc, "{{exhibit_a_equipment}}", font_size=10, space_after=Pt(12))


def build_exhibit_x(doc, entity_name="Customer"):
    doc.add_page_break()
    add_p(doc, "Confidential", font_size=9, space_after=Pt(18))
    add_p(doc, "EXHIBIT X", font_size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_p(doc, "Data Privacy Provisions", font_size=10, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(16))

    p = add_p(doc, font_size=10, line_spacing=1.15)
    r_head = p.add_run("1. Data Privacy and Security Requirements. ")
    r_head.bold = True
    r_head.underline = True
    r_head.font.name = "Times New Roman"
    r_head.font.size = Pt(10)
    r_body = p.add_run(
        f'To the extent {entity_name} accesses, collects, uses, stores, or otherwise processes any data in connection with this Agreement that would constitute '
        f'"Personal Information" or a similar data classification subject to additional requirements under any applicable data protection, data security, or privacy law '
        f'("Data Protection Law") {entity_name} shall comply with such Data Protection Law in addition to {entity_name}\'s other obligations regarding that data under this Agreement.'
    )
    r_body.font.name = "Times New Roman"
    r_body.font.size = Pt(10)


# =============================================================================
# 1. DIRECT AGREEMENT TEMPLATE - CUSTOMER OWNERSHIP
# =============================================================================
def create_direct_customer_ownership():
    doc = Document()
    
    # Page 1: Appendix B
    build_page_1_appendix_b(doc, is_indirect=False)

    # Page 2: Device Release Form
    add_p(doc, "Confidential", font_size=9, space_after=Pt(14))
    add_p(doc, "DEVICE RELEASE FORM", font_size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(14))

    add_p(doc, 'This DEVICE RELEASE FORM ("Form"), made as of this {DATE} ("Effective Date"), by and between', font_size=10, space_after=Pt(6))
    add_p(doc, 'Intervet India Private Limited, a company incorporated under the laws of India, having its registered office at Pune, Maharashtra ("MAH")', font_size=10, space_after=Pt(6))
    add_p(doc, 'and', font_size=10, space_after=Pt(6))
    add_p(doc, '{Customer Name}, a corporation duly organized and existing under the laws of India, having its principal place of business at {Customer Name},', font_size=10, space_after=Pt(2))
    add_p(doc, '{ADDRESS OF THE CUSTOMER COMPANY}', font_size=10, bold=False, space_after=Pt(10))

    add_p(doc, 'MAH and Customer, intending to be legally bound, hereby agree as follows:', font_size=10, space_after=Pt(8))

    clauses = [
        "1.  MAH shall deliver to the Customer the Device free of charge.",
        '2.  The Customer shall endeavor during the Term of this Form to purchase such quantities of Products as set forth in Exhibit A attached to this Form. At the end of each {calendar quarter and/or period} ("Period").',
        "3.  The Customer agrees to promptly inform the MAH of the deployment of any devices to its Customer. Such notification shall include detailed information regarding the location of each deployed device. This information is essential for the MAH to provide the necessary support as per the agreed terms. The Customer shall ensure that all deployment details are communicated to the MAH within 5 working days of the installation of the Device at its customer location.",
        "4.  The Customer agrees to hold harmless, release, and indemnify MAH in full of any and all liability, damage, loss, or harm related to the Device or arising in any way from the use or storage of the Device.",
        "5.  The Customer acknowledges that the Device was received in good working condition and is henceforth the sole and exclusive responsibility of the Customer and that MAH owes the Customer no further obligation, of any kind whatsoever, related to the Device."
    ]

    for cl in clauses:
        add_p(doc, cl, font_size=9.5, space_after=Pt(6))

    # Page 3: Clause 6 & Signatures
    doc.add_page_break()
    add_p(doc, "Confidential", font_size=9, space_after=Pt(14))
    add_p(doc, "6.  The Customer understands that MAH released ownership of the Device to the Customer and further agrees that the Customer shall hold harmless, release, and indemnify MAH in full of any and all liability, damage, loss, or harm related to the Device or arising in any way from use or storage of the Device.", font_size=9.5, space_after=Pt(14))

    build_signature_section(doc, left_party_name="{Distributor Name}", right_party_name="Intervet India Private Limited")

    # Page 4: Exhibit A
    build_exhibit_a(doc, "Products and Quantities")

    # Page 5: Exhibit X (reference image shows "Distributor" for Direct Customer Ownership)
    build_exhibit_x(doc, "Distributor")

    return doc


# =============================================================================
# 2. DIRECT AGREEMENT TEMPLATE - INNOJECT PRO
# =============================================================================
def create_direct_innoject_pro():
    doc = Document()
    
    # Page 1: Appendix B
    build_page_1_appendix_b(doc, is_indirect=False)

    # Page 2: Device Agreement
    add_p(doc, "Confidential", font_size=9, space_after=Pt(14))
    add_p(doc, "DEVICE AGREEMENT", font_size=11, bold=True, underline=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(14))

    add_p(doc, 'This DEVICE AGREEMENT ("Agreement"), made as of this {DATE} ("Effective Date"), by and between', font_size=10, space_after=Pt(6))
    add_p(doc, 'Intervet India Private Limited, a company incorporated under the laws of India, having its registered office at Pune, Maharashtra ("MAH")', font_size=10, space_after=Pt(6))
    add_p(doc, 'and', font_size=10, space_after=Pt(6))
    add_p(doc, '{Customer Name}, a corporation duty organized and existing under the laws of India, having its principal place of business at {Customer Name},', font_size=10, space_after=Pt(2))
    add_p(doc, '{ADDRESS OF THE CUSTOMER COMPANY}', font_size=10, bold=False, space_after=Pt(8))

    add_p(doc, 'WHEREAS,', font_size=10, bold=True, space_after=Pt(4))
    add_p(doc, '•  MAH is engaged in the business of developing, manufacturing, marketing and selling certain Products (as hereinafter defined);', font_size=9.5, space_after=Pt(3))
    add_p(doc, '•  MAH is willing to offer to Customer the opportunity to use the Device as complementary service and Customer wishes to accept such offer at the terms set herein.', font_size=9.5, space_after=Pt(6))

    add_p(doc, 'NOW, THEREFORE, in consideration of the promises and the mutual agreements, covenants and conditions set forth in this Agreement and other good and valuable consideration, the receipt and sufficiency of which are hereby acknowledged, MAH and Customer, intending to be legally bound, hereby agree as follows:', font_size=9.5, space_after=Pt(8))

    # Article 1
    p_art1 = add_p(doc, font_size=10, space_after=Pt(3))
    add_run_to_para(p_art1, "Article 1.    ", font_size=10)
    add_run_to_para(p_art1, "Definitions", font_size=10, underline=True)
    
    add_p(doc, 'For purposes of this Agreement, the following terms shall have the following meanings:', font_size=9.5, space_after=Pt(2))
    add_p(doc, '1.1  "Device" shall mean Innoject Pro Double Injection with Eye Drop.', font_size=9.5, space_after=Pt(2))
    add_p(doc, '1.2  "Party" shall mean MAH or Customer and "Parties" shall mean MAH and Customer.', font_size=9.5, space_after=Pt(2))
    add_p(doc, '1.3  "Products" shall mean any finished, packaged product, as set forth in Exhibit A, which is manufactured by or on behalf of MAH or any Affiliate of MAH in accordance with the standards, specifications, and formulae established by MAH or such Affiliate.', font_size=9.5, space_after=Pt(6))

    # Article 2
    p_art2 = add_p(doc, font_size=10, space_after=Pt(3))
    add_run_to_para(p_art2, "Article 2.    ", font_size=10)
    add_run_to_para(p_art2, "Usage of Device", font_size=10, underline=True)

    add_p(doc, '2.1  MAH shall provide the Device free of charge to the Customer as a complementary service. The MAH can provide these devices to its customers free of charge, complying with MAH\'s ownership rights as provided in this agreement.', font_size=9.5, space_after=Pt(3))
    add_p(doc, '2.2  The Customer agrees to promptly inform the MAH of the deployment of any devices to its customers. Such notification shall include detailed information regarding the location of each deployed device. This information is essential for the MAH to provide the necessary support as per the agreed terms. The Customer shall ensure that all deployment details are communicated to the MAH within 5 working days of the installation of the Device at its customer location.', font_size=9.5, space_after=Pt(6))

    # Page 3
    doc.add_page_break()
    add_p(doc, "Confidential", font_size=9, space_after=Pt(14))

    # Article 3
    p_art3 = add_p(doc, font_size=10, space_after=Pt(3))
    add_run_to_para(p_art3, "Article 3.    ", font_size=10)
    add_run_to_para(p_art3, "Transfer and Delivery", font_size=10, underline=True)

    add_p(doc, '3.1  MAH will deliver the Device at the address of Customer, as listed above. Delivery will take place as soon as possible after the signing of this Agreement by both Parties. Customer accept that Device was received in good working condition.', font_size=9.5, space_after=Pt(6))

    # Article 4
    p_art4 = add_p(doc, font_size=10, space_after=Pt(3))
    add_run_to_para(p_art4, "Article 4.    ", font_size=10)
    add_run_to_para(p_art4, "Ownership", font_size=10, underline=True)

    add_p(doc, '4.1  The Device remains the property of MAH. The Customer shall have no right or interest other than as a user of the Device. The Customer shall not sell, assign, sublet, pledge or otherwise dispose of the Device. The Customer shall ensure that the Device remains the property of MAH and will not fix the Device to anything so that it cannot be removed without causing damage to the Device.', font_size=9.5, space_after=Pt(3))
    add_p(doc, '4.2  If the Device is stolen, lost or damaged while in the possession of the Customer, the Customer shall at its cost and expenses replace the Device or pay damages to the reasonable satisfaction of MAH, unless otherwise agreed by the Parties. Customer will ensure the Device at its own expense and keep insured against damage by fire and theft.', font_size=9.5, space_after=Pt(3))
    add_p(doc, '4.3  The Customer shall not allow any third party to use the Device unless MAH agrees to this use in writing.', font_size=9.5, space_after=Pt(6))

    # Article 5
    p_art5 = add_p(doc, font_size=10, space_after=Pt(3))
    add_run_to_para(p_art5, "Article 5.    ", font_size=10)
    add_run_to_para(p_art5, "Use and Maintenance", font_size=10, underline=True)

    add_p(doc, '5.1  Customer undertakes to use or inform its customers to use the Device exclusively for the purposes and in accordance with the indications set forth in the instruction manual and/or as provided by MAH and in accordance with all applicable law.', font_size=9.5, space_after=Pt(3))
    add_p(doc, '5.2  Customer warrants that its customers use the Device throughout the Term of this Agreement with diligence and keep it in good working order and maintenance, except for ordinary wear and tear.', font_size=9.5, space_after=Pt(3))
    add_p(doc, '5.3  MAH will be responsible for the maintenance of the Device.', font_size=9.5, space_after=Pt(6))

    # Article 6
    p_art6_title = add_p(doc, font_size=10, space_after=Pt(3))
    add_run_to_para(p_art6_title, "6.    ", font_size=10)
    add_run_to_para(p_art6_title, "Warranty and Liability", font_size=10, underline=True)

    add_p(doc, '6.1  The Customer will be liable for all damages caused to person or property following the use of the Device and agrees to indemnify MAH from any claim for damage made by third parties related to the possession, safe-keeping or incorrect use of the Device.', font_size=9.5, space_after=Pt(3))
    add_p(doc, '6.2  MAH is not liable for the Device supplied for free. MAH will not be liable for any damages related to Device or caused by Customer\'s use of the Device.', font_size=9.5, space_after=Pt(6))

    # Article 7
    p_art7 = add_p(doc, font_size=10, space_after=Pt(3))
    add_run_to_para(p_art7, "Article 7.    ", font_size=10)
    add_run_to_para(p_art7, "Term and Termination", font_size=10, underline=True)

    add_p(doc, '7.1  This Agreement will be valid as of the Effective Date for a period of [number of years in letters, preferably no longer than 5 years], ([5]) year, thus ending on 13th of November, [2030].', font_size=9.5, space_after=Pt(6))

    # Page 4
    doc.add_page_break()
    add_p(doc, "Confidential", font_size=9, space_after=Pt(14))

    add_p(doc, '7.2  At termination or expiration of this Agreement, Customer will return the Device to MAH, at its own expense, on its own initiative or at the request of MAH within seven (7) days after receipt of such a request, whichever is earliest unless Parties agree otherwise in writing.', font_size=9.5, space_after=Pt(3))
    add_p(doc, '7.3  This Agreement may be terminated by MAH with immediate effect in case Customer did not purchase during the Period the agreed quantities of Products for that Period. In such event, at the sole discretion of MAH, Customer shall either pay to MAH the fair market value of the Device or Customer shall return the Device to MAH.', font_size=9.5, space_after=Pt(6))

    # Miscellaneous
    p_misc = add_p(doc, font_size=10, space_after=Pt(4))
    add_run_to_para(p_misc, "Miscellaneous", font_size=10, underline=True)

    # 7.4 Assignment
    p_74 = add_p(doc, font_size=9.5, space_after=Pt(4))
    add_run_to_para(p_74, "7.4    ", font_size=9.5)
    add_run_to_para(p_74, "Assignment", font_size=9.5, italic=True)
    add_p(doc, 'Customer shall not assign this Agreement or subcontract any of Customer\'s duties hereunder to any person, organization or other entity (including by operation of law, judicial process or otherwise) without the prior written consent of MAH, which consent may be withheld for any reason. MAH shall be entitled to assign this Agreement to any of its affiliates (including by operation of law, judicial process or otherwise) or any successor to its business or operations to which this Agreement relates without prior notice to or consent from Customer.', font_size=9.5, space_after=Pt(4))

    # 7.5 Entire Agreement
    p_75 = add_p(doc, font_size=9.5, space_after=Pt(4))
    add_run_to_para(p_75, "7.5    ", font_size=9.5)
    add_run_to_para(p_75, "Entire Agreement", font_size=9.5, italic=True)
    add_p(doc, 'This Agreement, including its Exhibits, represents and contains the full and complete understanding and agreement of the Parties with respect to the subject matter hereof and supersedes and replaces all prior and contemporaneous agreements, general conditions of either Party, understandings, statements, clauses and conditions, both oral and written, with respect to the transactions contemplated by this Agreement or which may be contained in any other form or document.', font_size=9.5, space_after=Pt(4))

    # 7.6 Amendment
    p_76 = add_p(doc, font_size=9.5, space_after=Pt(4))
    add_run_to_para(p_76, "7.6    ", font_size=9.5)
    add_run_to_para(p_76, "Amendment", font_size=9.5, italic=True)
    add_p(doc, 'Neither this Agreement nor any provision hereof may be amended, supplemented, waived or modified, except by a specific writing, entitled as an amendment and specifically referring to this Agreement, that is signed by an authorized officer of each Party.', font_size=9.5, space_after=Pt(4))

    # 7.7 Severability
    p_77 = add_p(doc, font_size=9.5, space_after=Pt(4))
    add_run_to_para(p_77, "7.7    ", font_size=9.5)
    add_run_to_para(p_77, "Severability", font_size=9.5, italic=True)
    add_p(doc, 'In the event that any one or more of the provisions in this Agreement shall, for any reason, be held to be invalid, illegal or unenforceable in any respect, such invalidity, illegality or unenforceability, shall not affect any other provisions of this Agreement and all other provisions shall remain in full force and effect.', font_size=9.5, space_after=Pt(4))

    # 7.8 Data Privacy
    p_78 = add_p(doc, font_size=9.5, space_after=Pt(4))
    add_run_to_para(p_78, "7.8    ", font_size=9.5)
    add_run_to_para(p_78, "Data Privacy", font_size=9.5, italic=True)
    add_p(doc, 'Each of Customer and MAH shall comply with the requirements of Exhibit X in connection with its obligations under this Agreement.', font_size=9.5, space_after=Pt(4))

    # 7.9 Governing Law
    p_79 = add_p(doc, font_size=9.5, space_after=Pt(4))
    add_run_to_para(p_79, "7.9    ", font_size=9.5)
    add_run_to_para(p_79, "Governing Law", font_size=9.5, italic=True)
    add_p(doc, 'This Agreement shall be construed and governed in accordance with the laws India, without giving effect to the conflict of laws, rules or principles thereof. All disputes', font_size=9.5, space_after=Pt(4))

    # Page 5: Dispute settlement continuation & Signatures
    doc.add_page_break()
    add_p(doc, "Confidential", font_size=9, space_after=Pt(14))
    add_p(doc, 'arising out, of or in connection with this Agreement, which cannot be settled amicably, shall be exclusively settled by the court of Pune, India.', font_size=9.5, space_after=Pt(14))

    build_signature_section(doc, left_party_name="{Distributor Name}", right_party_name="Intervet India Private Limited")

    # Page 6: Exhibit A
    build_exhibit_a(doc, "Proposed and Quantities")

    # Page 7: Exhibit X
    build_exhibit_x(doc, "Customer")

    return doc


# =============================================================================
# 3. INDIRECT AGREEMENT TEMPLATE - CUSTOMER OWNERSHIP
# =============================================================================
def create_indirect_customer_ownership():
    doc = Document()
    
    # Page 1: Appendix B
    build_page_1_appendix_b(doc, is_indirect=True)

    # Page 2: Device Release Form (Distributor)
    add_p(doc, "Confidential", font_size=9, space_after=Pt(14))

    p_title = add_p(doc, font_size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_run_to_para(p_title, "DEVICE RELEASE FORM", font_size=11, bold=True, underline=True)
    p_subtitle = add_p(doc, "(DISTRIBUTOR)", font_size=11, bold=True, underline=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(14))

    add_p(doc, 'This DEVICE RELEASE FORM ("Form"), made as of this {DATE} ("Effective Date"), by and between', font_size=10, space_after=Pt(6))
    add_p(doc, 'Intervet India Private Limited, a company incorporated under the laws of India, having its registered office at Pune, Maharashtra ("MAH")', font_size=10, space_after=Pt(6))
    add_p(doc, 'and', font_size=10, space_after=Pt(6))
    add_p(doc, '{Distributor Name}, a corporation duly organized and existing under the laws of {Country}, having its principal place of business at {Distributor Name},', font_size=10, space_after=Pt(2))
    add_p(doc, '{ADDRESS OF THE DISTRIBUTOR COMPANY}', font_size=10, bold=False, space_after=Pt(10))

    add_p(doc, 'MAH and Distributor, intending to be legally bound, hereby agree as follows:', font_size=10, space_after=Pt(8))

    clauses = [
        "1.  MAH shall deliver to the Distributor the Device free of charge.",
        '2.  The Distributor shall endeavor during the Term of this Form to purchase such quantities of Products as set forth in Exhibit A attached to this Form. At the end of each [calendar quarter and/or period] ("Period").',
        "3.  The Distributor agrees to promptly inform the MAH of the deployment of any devices to its customers. Such notification shall include detailed information regarding the location of each deployed device. This information is essential for the MAH to provide the necessary support as per the agreed terms. The Distributor shall ensure that all deployment details are communicated to the MAH within 5 working days of the installation of the Device at its customer location.",
        "4.  The Distributor agrees to hold harmless, release, and indemnify MAH in full of any and all liability, damage, loss, or harm related to the Device or arising in any way from the use or storage of the Device.",
        "5.  The Distributor acknowledges that the Device was received in good working condition and is henceforth the sole and exclusive responsibility of the Distributor and that MAH owes the Distributor no further obligation, of any kind whatsoever, related to the Device."
    ]

    for cl in clauses:
        add_p(doc, cl, font_size=9.5, space_after=Pt(6))

    # Page 3: Clause 6 & Signatures
    doc.add_page_break()
    add_p(doc, "Confidential", font_size=9, space_after=Pt(14))
    add_p(doc, "6.  The Customer understands that MAH released ownership of the Device to the Customer and further agrees that the Customer shall hold harmless, release, and indemnify MAH in full of any and all liability, damage, loss, or harm related to the Device or arising in any way from use or storage of the Device.", font_size=9.5, space_after=Pt(14))

    build_signature_section(doc, left_party_name="{Customer Name}", right_party_name="Intervet India Private Limited")

    # Page 4: Exhibit A
    build_exhibit_a(doc, "Products and Quantities")

    # Page 5: Exhibit X
    build_exhibit_x(doc, "Customer")

    return doc


# =============================================================================
# 4. INDIRECT AGREEMENT TEMPLATE - INNOJECT PRO
# =============================================================================
def create_indirect_innoject_pro():
    doc = Document()
    
    # Page 1: Appendix B
    build_page_1_appendix_b(doc, is_indirect=True)

    # Page 2: Device Agreement
    add_p(doc, "Confidential", font_size=9, space_after=Pt(14))
    add_p(doc, "DEVICE AGREEMENT", font_size=11, bold=True, underline=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(14))

    add_p(doc, 'This DEVICE AGREEMENT ("Agreement"), made as of this {DATE} ("Effective Date"), by and between', font_size=10, space_after=Pt(6))
    add_p(doc, 'Intervet India Private Limited, a company incorporated under the laws of India, having its registered office at Pune, Maharashtra ("MAH")', font_size=10, space_after=Pt(6))
    add_p(doc, 'and', font_size=10, space_after=Pt(6))
    add_p(doc, '{Distributor Name}, a corporation duly organized and existing under the laws of India, having its principal place of business at {Distributor Name},', font_size=10, space_after=Pt(2))
    add_p(doc, '{ADDRESS OF THE DISTRIBUTOR COMPANY}', font_size=10, bold=False, space_after=Pt(8))

    add_p(doc, 'WHEREAS,', font_size=10, bold=True, space_after=Pt(4))
    add_p(doc, '•  MAH is engaged in the business of developing, manufacturing, marketing and selling certain Products (as hereinafter defined);', font_size=9.5, space_after=Pt(3))
    add_p(doc, '•  MAH is willing to offer to Customer the opportunity to use the Device as complementary service and Customer wishes to accept such offer at the terms set herein.', font_size=9.5, space_after=Pt(6))

    add_p(doc, 'NOW, THEREFORE, in consideration of the promises and the mutual agreements, covenants and conditions set forth in this Agreement and other good and valuable consideration, the receipt and sufficiency of which are hereby acknowledged, MAH and Customer, intending to be legally bound, hereby agree as follows:', font_size=9.5, space_after=Pt(8))

    # Article 1
    p_art1 = add_p(doc, font_size=10, space_after=Pt(3))
    add_run_to_para(p_art1, "Article 1.    ", font_size=10)
    add_run_to_para(p_art1, "Definitions", font_size=10, underline=True)
    
    add_p(doc, 'For purposes of this Agreement, the following terms shall have the following meanings:', font_size=9.5, space_after=Pt(2))
    add_p(doc, '1.1  "Device" shall mean Innoject Pro Double Injection with Eye Drop.', font_size=9.5, space_after=Pt(2))
    add_p(doc, '1.2  "Party" shall mean MAH or Customer and "Parties" shall mean MAH and Customer.', font_size=9.5, space_after=Pt(2))
    add_p(doc, '1.3  "Products" shall mean any finished, packaged product, as set forth in Exhibit A, which is manufactured by or on behalf of MAH or any Affiliate of MAH in accordance with the standards, specifications, and formulae established by MAH or such Affiliate.', font_size=9.5, space_after=Pt(6))

    # Article 2
    p_art2 = add_p(doc, font_size=10, space_after=Pt(3))
    add_run_to_para(p_art2, "Article 2.    ", font_size=10)
    add_run_to_para(p_art2, "Usage of Device", font_size=10, underline=True)

    add_p(doc, '2.1  MAH shall provide the Device free of charge to the Customer as a complementary service. The MAH can provide these devices to its customers free of charge, complying with MAH\'s ownership rights as provided in this agreement.', font_size=9.5, space_after=Pt(3))
    add_p(doc, '2.2  The Customer agrees to promptly inform the MAH of the deployment of any devices to its customers. Such notification shall include detailed information regarding the location of each deployed device. This information is essential for the MAH to provide the necessary support as per the agreed terms. The Customer shall ensure that all deployment details are communicated to the MAH within 5 working days of the installation of the Device at its customer location.', font_size=9.5, space_after=Pt(6))

    # Page 3
    doc.add_page_break()
    add_p(doc, "Confidential", font_size=9, space_after=Pt(14))

    # Article 3
    p_art3 = add_p(doc, font_size=10, space_after=Pt(3))
    add_run_to_para(p_art3, "Article 3.    ", font_size=10)
    add_run_to_para(p_art3, "Transfer and Delivery", font_size=10, underline=True)

    add_p(doc, '3.1  MAH will deliver the Device at the address of Customer, as listed above. Delivery will take place as soon as possible after the signing of this Agreement by both Parties. Customer accept that Device was received in good working condition.', font_size=9.5, space_after=Pt(6))

    # Article 4
    p_art4 = add_p(doc, font_size=10, space_after=Pt(3))
    add_run_to_para(p_art4, "Article 4.    ", font_size=10)
    add_run_to_para(p_art4, "Ownership", font_size=10, underline=True)

    add_p(doc, '4.1  The Device remains the property of MAH. The Customer shall have no right or interest other than as a user of the Device. The Customer shall not sell, assign, sublet, pledge or otherwise dispose of the Device. The Customer shall ensure that the Device remains the property of MAH and will not fix the Device to anything so that it cannot be removed without causing damage to the Device.', font_size=9.5, space_after=Pt(3))
    add_p(doc, '4.2  If the Device is stolen, lost or damaged while in the possession of the Customer, the Customer shall at its cost and expenses replace the Device or pay damages to the reasonable satisfaction of MAH, unless otherwise agreed by the Parties. Customer will ensure the Device at its own expense and keep insured against damage by fire and theft.', font_size=9.5, space_after=Pt(3))
    add_p(doc, '4.3  The Customer shall not allow any third party to use the Device unless MAH agrees to this use in writing.', font_size=9.5, space_after=Pt(6))

    # Article 5
    p_art5 = add_p(doc, font_size=10, space_after=Pt(3))
    add_run_to_para(p_art5, "Article 5.    ", font_size=10)
    add_run_to_para(p_art5, "Use and Maintenance", font_size=10, underline=True)

    add_p(doc, '5.1  Customer undertakes to use or inform its customers to use the Device exclusively for the purposes and in accordance with the indications set forth in the instruction manual and/or as provided by MAH and in accordance with all applicable law.', font_size=9.5, space_after=Pt(3))
    add_p(doc, '5.2  Customer warrants that its customers use the Device throughout the Term of this Agreement with diligence and keep it in good working order and maintenance, except for ordinary wear and tear.', font_size=9.5, space_after=Pt(3))
    add_p(doc, '5.3  MAH will be responsible for the maintenance of the Device.', font_size=9.5, space_after=Pt(6))

    # Article 6
    p_art6_title = add_p(doc, font_size=10, space_after=Pt(3))
    add_run_to_para(p_art6_title, "6.    ", font_size=10)
    add_run_to_para(p_art6_title, "Warranty and Liability", font_size=10, underline=True)

    add_p(doc, '6.1  The Customer will be liable for all damages caused to person or property following the use of the Device and agrees to indemnify MAH from any claim for damage made by third parties related to the possession, safe-keeping or incorrect use of the Device.', font_size=9.5, space_after=Pt(3))
    add_p(doc, '6.2  MAH is not liable for the Device supplied for free. MAH will not be liable for any damages related to Device or caused by Customerr\'s use of the Device.', font_size=9.5, space_after=Pt(6))

    # Article 7
    p_art7 = add_p(doc, font_size=10, space_after=Pt(3))
    add_run_to_para(p_art7, "Article 7.    ", font_size=10)
    add_run_to_para(p_art7, "Term and Termination", font_size=10, underline=True)

    add_p(doc, '7.1  This Agreement will be valid as of the Effective Date for a period of [number of years in letters, preferably no longer than 5 years], ([5]) year, thus ending on 9th of February, [2031].', font_size=9.5, space_after=Pt(6))

    # Page 4
    doc.add_page_break()
    add_p(doc, "Confidential", font_size=9, space_after=Pt(14))

    add_p(doc, '7.2  At termination or expiration of this Agreement, Customer will return the Device to MAH, at its own expense, on its own initiative or at the request of MAH within seven (7) days after receipt of such a request, whichever is earliest unless Parties agree otherwise in writing.', font_size=9.5, space_after=Pt(3))
    add_p(doc, '7.3  This Agreement may be terminated by MAH with immediate effect in case Customer did not purchase during the Period the agreed quantities of Products for that Period. In such event, at the sole discretion of MAH, Customer shall either pay to MAH the fair market value of the Device or Customer shall return the Device to MAH.', font_size=9.5, space_after=Pt(6))

    # Miscellaneous
    p_misc = add_p(doc, font_size=10, space_after=Pt(4))
    add_run_to_para(p_misc, "Miscellaneous", font_size=10, underline=True)

    # 7.4 Assignment
    p_74 = add_p(doc, font_size=9.5, space_after=Pt(4))
    add_run_to_para(p_74, "7.4    ", font_size=9.5)
    add_run_to_para(p_74, "Assignment", font_size=9.5, italic=True)
    add_p(doc, 'Customer shall not assign this Agreement or subcontract any of Customer\'s duties hereunder to any person, organization or other entity (including by operation of law, judicial process or otherwise) without the prior written consent of MAH, which consent may be withheld for any reason. MAH shall be entitled to assign this Agreement to any of its affiliates (including by operation of law, judicial process or otherwise) or any successor to its business or operations to which this Agreement relates without prior notice to or consent from Customer.', font_size=9.5, space_after=Pt(4))

    # 7.5 Entire Agreement
    p_75 = add_p(doc, font_size=9.5, space_after=Pt(4))
    add_run_to_para(p_75, "7.5    ", font_size=9.5)
    add_run_to_para(p_75, "Entire Agreement", font_size=9.5, italic=True)
    add_p(doc, 'This Agreement, including its Exhibits, represents and contains the full and complete understanding and agreement of the Parties with respect to the subject matter hereof and supersedes and replaces all prior and contemporaneous agreements, general conditions of either Party, understandings, statements, clauses and conditions, both oral and written, with respect to the transactions contemplated by this Agreement or which may be contained in any other form or document.', font_size=9.5, space_after=Pt(4))

    # 7.6 Amendment
    p_76 = add_p(doc, font_size=9.5, space_after=Pt(4))
    add_run_to_para(p_76, "7.6    ", font_size=9.5)
    add_run_to_para(p_76, "Amendment", font_size=9.5, italic=True)
    add_p(doc, 'Neither this Agreement nor any provision hereof may be amended, supplemented, waived or modified, except by a specific writing, entitled as an amendment and specifically referring to this Agreement, that is signed by an authorized officer of each Party.', font_size=9.5, space_after=Pt(4))

    # 7.7 Severability
    p_77 = add_p(doc, font_size=9.5, space_after=Pt(4))
    add_run_to_para(p_77, "7.7    ", font_size=9.5)
    add_run_to_para(p_77, "Severability", font_size=9.5, italic=True)
    add_p(doc, 'In the event that any one or more of the provisions in this Agreement shall, for any reason, be held to be invalid, illegal or unenforceable in any respect, such invalidity, illegality or unenforceability, shall not affect any other provisions of this Agreement and all other provisions shall remain in full force and effect.', font_size=9.5, space_after=Pt(4))

    # 7.8 Data Privacy
    p_78 = add_p(doc, font_size=9.5, space_after=Pt(4))
    add_run_to_para(p_78, "7.8    ", font_size=9.5)
    add_run_to_para(p_78, "Data Privacy", font_size=9.5, italic=True)
    add_p(doc, 'Each of Customer and MAH shall comply with the requirements of Exhibit X in connection with its obligations under this Agreement.', font_size=9.5, space_after=Pt(4))

    # 7.9 Governing Law
    p_79 = add_p(doc, font_size=9.5, space_after=Pt(4))
    add_run_to_para(p_79, "7.9    ", font_size=9.5)
    add_run_to_para(p_79, "Governing Law", font_size=9.5, italic=True)
    add_p(doc, 'This Agreement shall be construed and governed in accordance with the laws India, without giving effect to the conflict of laws, rules or principles thereof. All disputes', font_size=9.5, space_after=Pt(4))

    # Page 5: Dispute settlement continuation & Signatures
    doc.add_page_break()
    add_p(doc, "Confidential", font_size=9, space_after=Pt(14))
    add_p(doc, 'arising out, of or in connection with this Agreement, which cannot be settled amicably, shall be exclusively settled by the court of Pune, India.', font_size=9.5, space_after=Pt(14))

    build_signature_section(doc, left_party_name="{Customer Name}", right_party_name="Intervet India Private Limited")

    # Page 6: Exhibit A
    build_exhibit_a(doc, "Proposed and Quantities")

    # Page 7: Exhibit X
    build_exhibit_x(doc, "Customer")

    return doc


# =============================================================================
# MAIN EXECUTOR
# =============================================================================
TEMPLATES = {
    'Direct Agreement Template-Customer ownership': create_direct_customer_ownership,
    'Direct Agreement Template-Innoject Pro': create_direct_innoject_pro,
    'Indirect Agreement Template-Customer Ownership': create_indirect_customer_ownership,
    'Indirect Agreement Template-Innoject Pro': create_indirect_innoject_pro,
}


def generate_all_templates():
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    for name, creator_fn in TEMPLATES.items():
        filepath = TEMPLATES_DIR / f"{name}.docx"
        doc = creator_fn()
        doc.save(filepath)
        print(f"  [OK] Created matching template: {filepath.name}")


if __name__ == '__main__':
    print("Generating exact visual-matching agreement templates from reference images...")
    generate_all_templates()
    print("All 4 exact templates created successfully.")
