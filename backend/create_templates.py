from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from config import TEMPLATES_DIR, AGREEMENT_TYPES

def create_template(title: str, filename: str):
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run("MSD AGREEMENT")
    header_run.bold = True
    header_run.font.size = Pt(16)
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title.upper())
    title_run.bold = True
    title_run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    # Parties
    doc.add_heading("1. PARTIES", level=1)
    doc.add_paragraph("This Agreement is made on {{agreement_start_date}} between MSD and:")
    doc.add_paragraph("Company Name: {{company_name}}")
    doc.add_paragraph("Address: {{customer_address}}")
    doc.add_paragraph("Contact Person: {{contact_person_name}} ({{contact_person_designation}})")
    doc.add_paragraph("Email: {{contact_person_email}}")
    doc.add_paragraph("Phone: {{contact_person_phone}}")
    
    # Device Details
    doc.add_heading("2. DEVICE DETAILS", level=1)
    doc.add_paragraph("Device Name: {{device_name}}")
    doc.add_paragraph("Serial Number: {{device_serial_number}}")
    doc.add_paragraph("Territory: {{territory}}")
    
    # Terms
    doc.add_heading("3. TERMS", level=1)
    doc.add_paragraph("Agreement Value: {{agreement_value}}")
    doc.add_paragraph("Start Date: {{agreement_start_date}}")
    doc.add_paragraph("End Date: {{agreement_end_date}}")
    doc.add_paragraph("Duration: {{agreement_duration}}")
    
    doc.add_paragraph("\n" * 2)
    
    # Signatures
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "For MSD:"
    table.cell(0, 1).text = "For {{company_name}}:"
    table.cell(1, 0).text = "{{msd_signature}}"
    table.cell(1, 1).text = "{{customer_signature}}"
    
    doc.save(TEMPLATES_DIR / filename)

def generate_all_templates():
    for category, types in AGREEMENT_TYPES.items():
        for t in types:
            filename = f"{t}.docx"
            path = TEMPLATES_DIR / filename
            if not path.exists():
                create_template(t, filename)

if __name__ == "__main__":
    generate_all_templates()
    print("Templates created successfully.")
