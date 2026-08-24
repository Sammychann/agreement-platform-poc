import uuid
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from PIL import Image

from config import TEMPLATES_DIR, GENERATED_DIR


def format_day_suffix(day: int) -> str:
    if 11 <= (day % 100) <= 13:
        suffix = 'th'
    else:
        suffix = {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    return f"{day}{suffix}"


def format_agreement_date(date_str: str) -> str:
    """Format date to '28th May, 2026' style."""
    if not date_str:
        return ""
    for fmt in ('%Y-%m-%d', '%d/%m/%Y', '%d-%m-%Y', '%Y/%m/%d', '%Y-%m-%d %H:%M:%S'):
        try:
            dt = datetime.strptime(date_str.strip()[:10], fmt[:len(date_str.strip()[:10])])
            day_str = format_day_suffix(dt.day)
            month_name = dt.strftime('%B')
            year = dt.year
            return f"{day_str} {month_name}, {year}"
        except Exception:
            pass
    return str(date_str)


def format_signature_date(date_str: str) -> str:
    """Format date to 'DD/MM/YYYY' style."""
    if not date_str:
        return ""
    for fmt in ('%Y-%m-%d', '%d/%m/%Y', '%d-%m-%Y', '%Y/%m/%d'):
        try:
            dt = datetime.strptime(date_str.strip()[:10], fmt[:len(date_str.strip()[:10])])
            return dt.strftime('%d/%m/%Y')
        except Exception:
            pass
    return str(date_str)


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


class AgreementGenerator:

    def _replace_text_in_paragraph(self, paragraph, placeholders: Dict[str, str]):
        for key, val in placeholders.items():
            if key in paragraph.text:
                full_text = paragraph.text
                if key in full_text:
                    new_text = full_text.replace(key, str(val))
                    # Clear runs and set new text on first run or paragraph
                    if paragraph.runs:
                        paragraph.runs[0].text = new_text
                        for r in paragraph.runs[1:]:
                            r.text = ""
                    else:
                        paragraph.text = new_text

    def _replace_placeholders(self, doc: Document, placeholders: Dict[str, str]):
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            self._replace_text_in_paragraph(paragraph, placeholders)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_text_in_paragraph(paragraph, placeholders)

        # Replace in headers and footers
        for section in doc.sections:
            for header_para in section.header.paragraphs:
                self._replace_text_in_paragraph(header_para, placeholders)
            for footer_para in section.footer.paragraphs:
                self._replace_text_in_paragraph(footer_para, placeholders)

    def _normalize_image(self, img_path_str: str) -> Optional[str]:
        if not img_path_str:
            return None
        img_path = Path(img_path_str)
        if not img_path.exists():
            return None
        try:
            with Image.open(img_path) as im:
                clean_path = img_path.parent / f"norm_{img_path.name}"
                im.convert("RGBA").save(clean_path, format="PNG")
                return str(clean_path)
        except Exception:
            return str(img_path)

    def _embed_signatures(self, doc: Document, customer_sig: Optional[str], intervet_sig: Optional[str]):
        clean_cust_sig = self._normalize_image(customer_sig) if customer_sig else None
        clean_intervet_sig = self._normalize_image(intervet_sig) if intervet_sig else None

        # Replace in tables (signatures are inside borderless tables for layout)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if '{{customer_signature}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{customer_signature}}', '')
                            if clean_cust_sig and Path(clean_cust_sig).exists():
                                try:
                                    run = paragraph.add_run()
                                    run.add_picture(clean_cust_sig, width=Inches(1.8))
                                except Exception:
                                    pass
                        if '{{intervet_signature}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{intervet_signature}}', '')
                            if clean_intervet_sig and Path(clean_intervet_sig).exists():
                                try:
                                    run = paragraph.add_run()
                                    run.add_picture(clean_intervet_sig, width=Inches(1.8))
                                except Exception:
                                    pass

        # Also replace in normal paragraphs if any
        for paragraph in doc.paragraphs:
            if '{{customer_signature}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{customer_signature}}', '')
                if clean_cust_sig and Path(clean_cust_sig).exists():
                    try:
                        run = paragraph.add_run()
                        run.add_picture(clean_cust_sig, width=Inches(1.8))
                    except Exception:
                        pass
            if '{{intervet_signature}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{intervet_signature}}', '')
                if clean_intervet_sig and Path(clean_intervet_sig).exists():
                    try:
                        run = paragraph.add_run()
                        run.add_picture(clean_intervet_sig, width=Inches(1.8))
                    except Exception:
                        pass

    def _build_equipment_table(self, doc: Document, equipment_list: List[Dict[str, Any]], target_p):
        """Creates a formatted equipment table right where target_p is located."""
        # Insert table after target_p
        num_rows = len(equipment_list) + 1
        table = doc.add_table(rows=num_rows, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Widths
        for row in table.rows:
            row.cells[0].width = Inches(4.5)
            row.cells[1].width = Inches(2.0)

        # Header
        headers = ['Equipment Name', 'Quantity']
        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = h
            set_cell_shading(cell, 'E6F3F2')
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'

        # Rows
        for idx, item in enumerate(equipment_list):
            eq_name = str(item.get('equipment_name', '') if isinstance(item, dict) else getattr(item, 'equipment_name', ''))
            qty = str(item.get('quantity', '') if isinstance(item, dict) else getattr(item, 'quantity', ''))
            
            cell_0 = table.cell(idx + 1, 0)
            cell_0.text = eq_name
            for p in cell_0.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'

            cell_1 = table.cell(idx + 1, 1)
            cell_1.text = qty
            for p in cell_1.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'

        # Move table xml right after target_p xml
        target_p._p.addnext(table._tbl)
        # Remove target_p
        target_p._p.getparent().remove(target_p._p)

    def _build_appendix_b_equipment(self, doc: Document, equipment_list: List[Dict[str, Any]], target_p):
        """Replace the inline Appendix B equipment table with actual equipment data.
        
        This replaces the placeholder equipment rows in the Appendix B table
        (Description/Quantity/FMV format) with the actual equipment items.
        """
        # Build a new equipment table matching the Appendix B format (3 cols: Description, Quantity, FMV)
        num_rows = len(equipment_list) + 2  # header + data rows + empty row
        table = doc.add_table(rows=num_rows, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        for row in table.rows:
            row.cells[0].width = Inches(2.5)
            row.cells[1].width = Inches(1.5)
            row.cells[2].width = Inches(1.5)

        # Header row
        for i, h in enumerate(['Description', 'Quantity', 'FMV']):
            cell = table.cell(0, i)
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.underline = i < 2  # Description and Quantity are underlined
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'

        # Equipment data rows
        for idx, item in enumerate(equipment_list):
            eq_name = str(item.get('equipment_name', '') if isinstance(item, dict) else getattr(item, 'equipment_name', ''))
            qty = str(item.get('quantity', '') if isinstance(item, dict) else getattr(item, 'quantity', ''))
            
            table.cell(idx + 1, 0).text = eq_name
            table.cell(idx + 1, 1).text = qty
            table.cell(idx + 1, 2).text = ''
            
            for col in range(3):
                for p in table.cell(idx + 1, col).paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'

        # Empty trailing row
        last_row = num_rows - 1
        for col in range(3):
            table.cell(last_row, col).text = ''

        # Move table to position and remove the marker paragraph
        target_p._p.addnext(table._tbl)
        target_p._p.getparent().remove(target_p._p)

    def _replace_equipment_markers(self, doc: Document, equipment_list: List[Dict[str, Any]]):
        """Find {{appendix_b_equipment_table}} and {{exhibit_a_equipment}} and insert dynamic tables."""
        for p in list(doc.paragraphs):
            if '{{appendix_b_equipment_table}}' in p.text:
                self._build_appendix_b_equipment(doc, equipment_list, p)
            elif '{{exhibit_a_equipment}}' in p.text:
                self._build_equipment_table(doc, equipment_list, p)
        
        # Also check inside table cells for markers (Appendix B marker may be inside a table)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in list(cell.paragraphs):
                        if '{{appendix_b_equipment_table}}' in p.text:
                            # Just remove the marker from within the table cell
                            p.text = p.text.replace('{{appendix_b_equipment_table}}', '')
                        elif '{{exhibit_a_equipment}}' in p.text:
                            p.text = p.text.replace('{{exhibit_a_equipment}}', '')

    def _replace_equipment_in_existing_tables(self, doc: Document, equipment_list: List[Dict[str, Any]]):
        """Replace placeholder equipment rows in existing Appendix B tables.
        
        Finds cells containing '{Equipment Name}' and '{quantity}' in the Appendix B 
        equipment table (Description/Quantity/FMV) and replaces with actual data.
        """
        for table in doc.tables:
            eq_idx = 0
            for row in table.rows:
                cells = row.cells
                has_eq_placeholder = False
                for cell in cells:
                    if '{Equipment Name}' in cell.text or '{quantity}' in cell.text:
                        has_eq_placeholder = True
                        break
                
                if has_eq_placeholder and eq_idx < len(equipment_list):
                    item = equipment_list[eq_idx]
                    eq_name = str(item.get('equipment_name', '') if isinstance(item, dict) else getattr(item, 'equipment_name', ''))
                    qty = str(item.get('quantity', '') if isinstance(item, dict) else getattr(item, 'quantity', ''))
                    
                    for cell in cells:
                        for paragraph in cell.paragraphs:
                            if '{Equipment Name}' in paragraph.text:
                                new_text = paragraph.text.replace('{Equipment Name}', eq_name)
                                if paragraph.runs:
                                    paragraph.runs[0].text = new_text
                                    for r in paragraph.runs[1:]:
                                        r.text = ""
                                else:
                                    paragraph.text = new_text
                            if '{quantity}' in paragraph.text:
                                new_text = paragraph.text.replace('{quantity}', qty)
                                if paragraph.runs:
                                    paragraph.runs[0].text = new_text
                                    for r in paragraph.runs[1:]:
                                        r.text = ""
                                else:
                                    paragraph.text = new_text
                    eq_idx += 1

    def generate_agreement(self, entry_id: str, form_data: Dict[str, Any], agreement_type: str, 
                           customer_signature_path: Optional[str], intervet_signature_path: str) -> str:
        agreement_id = str(uuid.uuid4())
        
        template_name = f"{agreement_type}.docx"
        template_path = TEMPLATES_DIR / template_name
        
        if not template_path.exists():
            raise FileNotFoundError(f"Template {template_name} not found at {template_path}")
            
        doc = Document(template_path)
        
        formatted_date = format_agreement_date(str(form_data.get('date', '')))
        receiver_date = format_signature_date(str(form_data.get('receiver_date', '')))
        intervet_date = format_signature_date(str(form_data.get('intervet_date', '')))

        customer_name = str(form_data.get('customer_name', ''))
        distributor_name = str(form_data.get('distributor_name', '') or '')
        location = str(form_data.get('location', ''))
        address = str(form_data.get('address', ''))
        initiator_info = str(form_data.get('initiator_name_and_date', ''))
        manager_info = str(form_data.get('manager_name_and_date', ''))
        receiver_name = str(form_data.get('receiver_name', ''))
        receiver_title = str(form_data.get('receiver_title', ''))
        intervet_name = str(form_data.get('intervet_name', ''))
        intervet_title = str(form_data.get('intervet_title', ''))

        # Prepare equipment list
        raw_equipment = form_data.get('equipment', [])
        equipment_list = []
        if isinstance(raw_equipment, list):
            for item in raw_equipment:
                if isinstance(item, dict):
                    equipment_list.append(item)
                elif hasattr(item, 'dict'):
                    equipment_list.append(item.dict())
                elif hasattr(item, 'equipment_name'):
                    equipment_list.append({'equipment_name': item.equipment_name, 'quantity': item.quantity})

        # Build placeholders mapping - uses exact placeholder strings from templates
        placeholders = {
            # Core entity names
            '{Customer Name}': customer_name,
            '{Distributor Name}': distributor_name if distributor_name else customer_name,
            '{{Customer Name}}': customer_name,
            
            # Location
            '{Location}': location,
            
            # Date (formatted as "28th May, 2026")
            '{DATE}': formatted_date,
            
            # Addresses
            '{ADDRESS OF THE CUSTOMER COMPANY}': address,
            '{ADDRESS OF THE DISTRIBUTOR COMPANY}': address,
            
            # Internal approval
            '{Initiator Name and Date}': initiator_info,
            '{Manager Name and Date}': manager_info,
            
            # Customer/Receiver signature fields
            '{Receiver Name}': receiver_name,
            '{Title of receiver}': receiver_title,
            
            # Intervet signature fields
            '{Name}': intervet_name,
            '{Title}': intervet_title,
            
            # Country placeholder for indirect
            '{Country}': 'India',
        }
        
        # 1. Replace text placeholders throughout document
        self._replace_placeholders(doc, placeholders)

        # 2. Replace equipment placeholders in existing Appendix B tables
        self._replace_equipment_in_existing_tables(doc, equipment_list)

        # 3. Insert dynamic equipment tables for markers
        self._replace_equipment_markers(doc, equipment_list)

        # 4. Embed signatures
        self._embed_signatures(doc, customer_signature_path, intervet_signature_path)
        
        output_filename = f"{entry_id}_{agreement_type}_{agreement_id}.docx"
        output_path = GENERATED_DIR / output_filename
        doc.save(output_path)
        
        return agreement_id

    def edit_agreement(self, agreement_id: str, updated_fields: Dict[str, Any], entry: Dict[str, Any]) -> str:
        entry_id = entry.get('entry_id')
        agreement_type = updated_fields.get('agreement_type') or entry.get('agreement_type')
        cust_sig = entry.get('customer_signature_path')
        intervet_sig = entry.get('intervet_signature_path')
        
        form_data = {**entry, **updated_fields}
        new_agreement_id = self.generate_agreement(entry_id, form_data, agreement_type, cust_sig, intervet_sig)
        return new_agreement_id

    def get_agreement_path(self, agreement_id: str) -> Optional[Path]:
        for file in GENERATED_DIR.glob(f"*_{agreement_id}.docx"):
            return file
        return None

agreement_generator = AgreementGenerator()
