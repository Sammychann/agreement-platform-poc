"""
=============================================================================
AGREEMENT GENERATION MODULE (agreement_generation.py)
=============================================================================
This module is responsible for:
1. Validating that all mandatory dynamic fields have been provided.
2. Loading the designated Word (.docx) template from 'templates/'.
3. Populating ONLY the dynamic placeholders (e.g., {{ customer_name }}).
4. Preserving 100% of all fixed clauses, headings, formatting, and tables.
5. Saving the completed Word document to the 'output/' folder.
6. Generating a matching PDF version of the agreement.
7. Extracting structured preview data for the web interface.

KEY PRINCIPLE:
The Word template is the SOURCE OF TRUTH.
No contract text or clauses are hardcoded in Python.
=============================================================================
"""

import os
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Tuple, List, Optional

# Word template rendering engine
import docx
from docxtpl import DocxTemplate

# Local configuration
from config import (
    AGREEMENT_TYPES,
    TEMPLATES_DIR,
    OUTPUT_DIR,
    get_agreement_config
)


def validate_input_data(agreement_type_id: str, form_data: Dict[str, Any]) -> Tuple[bool, List[str]]:
    """
    Validates user form data against the required fields configured for the agreement type.

    Args:
        agreement_type_id: The ID of the agreement type (e.g., 'type_1')
        form_data: Dictionary of input values from the user/form

    Returns:
        Tuple of (is_valid: bool, error_messages: list of strings)
    """
    config = get_agreement_config(agreement_type_id)
    errors = []

    for field in config.get("fields", []):
        key = field["key"]
        label = field["label"]
        is_required = field.get("required", False)

        value = form_data.get(key)
        # Check if required field is empty or missing
        if is_required:
            if value is None or str(value).strip() == "":
                errors.append(f"'{label}' is a mandatory field and cannot be empty.")

    return (len(errors) == 0, errors)


def find_template_file(template_filename: str, alternate_name: Optional[str] = None) -> Path:
    """
    Locates the template file in the templates directory.
    Checks primary template_filename first, then alternate_name if provided.
    """
    primary_path = TEMPLATES_DIR / template_filename
    if primary_path.exists():
        return primary_path

    if alternate_name:
        alt_path = TEMPLATES_DIR / alternate_name
        if alt_path.exists():
            return alt_path

    raise FileNotFoundError(
        f"Word template not found. Looked for '{template_filename}' "
        f"and '{alternate_name}' in {TEMPLATES_DIR}"
    )


def sanitize_filename(name: str) -> str:
    """
    Cleans a string to make it safe for use in filenames.
    """
    # Replace spaces with underscores and remove non-alphanumeric characters
    cleaned = re.sub(r'[^a-zA-Z0-9_\-]', '_', name.strip())
    # Collapse multiple consecutive underscores
    cleaned = re.sub(r'_+', '_', cleaned)
    return cleaned.strip('_')[:50]


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    """
    Converts a Word (.docx) document to PDF.
    
    Approach:
    1. First tries native Microsoft Word conversion on Windows using 'docx2pdf'.
    2. If Microsoft Word is not installed or raises an error, gracefully falls back
       to 'reportlab' to create a clean PDF layout from the document text and tables.
    
    Returns:
        True if PDF conversion was successful, False otherwise.
    """
    # Attempt 1: Native Windows Microsoft Word COM conversion via docx2pdf
    try:
        from docx2pdf import convert
        # convert(input_path, output_path)
        convert(str(docx_path), str(pdf_path))
        if pdf_path.exists() and pdf_path.stat().st_size > 0:
            return True
    except Exception as e:
        print(f"[PDF Conversion Note] Native Word COM conversion unavailable ({e}). Using ReportLab engine.")

    # Attempt 2: High-fidelity ReportLab PDF Generator fallback
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch

        # Read the generated docx to extract text & tables
        doc = docx.Document(str(docx_path))
        pdf = SimpleDocTemplate(
            str(pdf_path),
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54
        )

        styles = getSampleStyleSheet()
        normal_style = styles["Normal"]
        normal_style.fontSize = 10
        normal_style.leading = 14
        normal_style.textColor = colors.HexColor("#1e293b")

        title_style = ParagraphStyle(
            'DocTitle',
            parent=styles['Heading1'],
            fontSize=16,
            leading=20,
            textColor=colors.HexColor("#0f172a"),
            alignment=1, # Center
            spaceAfter=12
        )

        heading_style = ParagraphStyle(
            'SectionHeading',
            parent=styles['Heading2'],
            fontSize=12,
            leading=16,
            textColor=colors.HexColor("#0f766e"), # Professional Teal
            spaceBefore=10,
            spaceAfter=6,
            keepWithNext=True
        )

        story = []

        # Process paragraphs
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                story.append(Spacer(1, 6))
                continue

            # Detect title/headings
            if p.style.name.startswith("Heading 1") or ("AGREEMENT" in text.upper() and len(text) < 60):
                story.append(Paragraph(f"<b>{text}</b>", title_style))
                story.append(Spacer(1, 4))
            elif p.style.name.startswith("Heading") or re.match(r'^\d+\.', text):
                story.append(Paragraph(f"<b>{text}</b>", heading_style))
            else:
                story.append(Paragraph(text, normal_style))
                story.append(Spacer(1, 4))

        # Process tables
        for t in doc.tables:
            table_data = []
            for row in t.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace('\n', '<br/>')
                    row_cells.append(Paragraph(cell_text, normal_style))
                table_data.append(row_cells)

            if table_data:
                col_count = len(table_data[0])
                avail_width = 504 # letter width 612 - 108 margin
                col_width = avail_width / col_count

                pdf_table = Table(table_data, colWidths=[col_width] * col_count)
                pdf_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#f1f5f9")),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor("#0f172a")),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('INNERGRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                    ('BOX', (0, 0), (-1, -1), 1, colors.HexColor("#94a3b8")),
                    ('TOPPADDING', (0, 0), (-1, -1), 5),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                ]))
                story.append(Spacer(1, 8))
                story.append(pdf_table)
                story.append(Spacer(1, 10))

        pdf.build(story)
        return pdf_path.exists() and pdf_path.stat().st_size > 0

    except Exception as e:
        print(f"[PDF Fallback Error]: {e}")
        return False


def get_agreement_preview(docx_path: Path) -> Dict[str, Any]:
    """
    Extracts structured content from a generated Word document for in-browser preview.
    
    Returns:
        Dictionary containing:
        - title: Main title if found
        - sections: List of paragraphs and tables formatted for display
    """
    doc = docx.Document(str(docx_path))
    content_blocks = []
    title = "Commercial Agreement Preview"

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        
        # Check for headings or main title
        if p.style.name.startswith("Heading 1") or ("AGREEMENT" in text.upper() and len(text) < 70):
            title = text
            content_blocks.append({"type": "title", "text": text})
        elif p.style.name.startswith("Heading") or re.match(r'^\d+\.', text):
            content_blocks.append({"type": "heading", "text": text})
        else:
            content_blocks.append({"type": "paragraph", "text": text})

    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            row_vals = [cell.text.strip() for cell in row.cells]
            table_rows.append(row_vals)
        if table_rows:
            content_blocks.append({"type": "table", "rows": table_rows})

    return {
        "title": title,
        "blocks": content_blocks,
        "total_paragraphs": len(doc.paragraphs),
        "total_tables": len(doc.tables)
    }


def generate_agreement(
    agreement_type_id: str,
    form_data: Dict[str, Any],
    custom_filename: Optional[str] = None
) -> Dict[str, Any]:
    """
    Main entry point for Agreement Generation.

    Steps:
    1. Validates input data against mandatory fields.
    2. Identifies and loads the appropriate Word template (.docx).
    3. Populates ONLY dynamic fields into template context.
    4. Saves the rendered document as .docx in the output/ folder.
    5. Converts the generated .docx to .pdf.
    6. Extracts preview data for UI rendering.

    Args:
        agreement_type_id: ID of the agreement type (e.g., 'type_1')
        form_data: User-supplied dictionary of field values
        custom_filename: Optional custom filename prefix

    Returns:
        Dictionary with generation details, paths, and preview data.
    """
    # Step 1: Validate input
    is_valid, errors = validate_input_data(agreement_type_id, form_data)
    if not is_valid:
        return {
            "status": "error",
            "message": "Mandatory fields are missing or invalid.",
            "errors": errors
        }

    # Step 2: Retrieve configuration & template
    config = get_agreement_config(agreement_type_id)
    template_file = config.get("template_file")
    alt_name = config.get("alternate_name")

    template_path = find_template_file(template_file, alt_name)

    # Step 3: Load Word template via docxtpl
    # docxtpl preserves all formatting, styles, tables, and fixed clauses
    doc = DocxTemplate(str(template_path))

    # Sanitize and prepare context for Jinja2 template rendering
    context = {}
    for k, v in form_data.items():
        if v is None:
            context[k] = ""
        else:
            context[k] = str(v).strip()

    # Step 4: Render dynamic fields into template with autoescape for safe XML handling
    doc.render(context, autoescape=True)

    # Step 5: Save generated .docx document
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    customer_raw = form_data.get("customer_name", "Customer")
    safe_customer = sanitize_filename(customer_raw)
    
    if custom_filename:
        base_name = f"{sanitize_filename(custom_filename)}_{timestamp}"
    else:
        base_name = f"{safe_customer}_{agreement_type_id}_{timestamp}"

    docx_filename = f"{base_name}.docx"
    pdf_filename = f"{base_name}.pdf"

    docx_path = OUTPUT_DIR / docx_filename
    pdf_path = OUTPUT_DIR / pdf_filename

    doc.save(str(docx_path))

    # Step 6: Generate matching PDF
    pdf_success = convert_docx_to_pdf(docx_path, pdf_path)

    # Step 7: Build structured preview
    preview_data = get_agreement_preview(docx_path)

    return {
        "status": "success",
        "agreement_type_id": agreement_type_id,
        "agreement_name": config["name"],
        "template_used": template_path.name,
        "docx_filename": docx_filename,
        "pdf_filename": pdf_filename if pdf_success else None,
        "pdf_available": pdf_success,
        "docx_path": str(docx_path),
        "pdf_path": str(pdf_path) if pdf_success else None,
        "preview": preview_data,
        "populated_fields": context,
        "generated_at": datetime.now().isoformat()
    }
