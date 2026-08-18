import zipfile
import openpyxl
from docx import Document
from pathlib import Path
import os

base_dir = Path(r"c:\Users\Admin\Desktop\agreement poc\temp_dummy_data")
zip_output = Path(r"c:\Users\Admin\Desktop\agreement poc\sample_validation_data.zip")

# Clean temp directory
if base_dir.exists():
    for root, dirs, files in os.walk(base_dir, topdown=False):
        for f in files: os.remove(os.path.join(root, f))
        for d in dirs: os.rmdir(os.path.join(root, d))
    base_dir.rmdir()

# Structure: Monthly ZIP -> January -> 2 Companies (Apollo Pharmacy & Fortis Healthcare)
co1_dir = base_dir / "January" / "Apollo Pharmacy"
co2_dir = base_dir / "January" / "Fortis Healthcare"
co1_dir.mkdir(parents=True, exist_ok=True)
co2_dir.mkdir(parents=True, exist_ok=True)

# --- COMPANY 1: Apollo Pharmacy (COMPLETE & PASSING) ---

# 1. Agreement (.docx)
doc1 = Document()
doc1.add_heading("AGREEMENT", level=1)
doc1.add_paragraph("Company Name: Apollo Pharmacy")
doc1.add_paragraph("Agreement Value: INR 5,00,000")
doc1.add_paragraph("Start Date: 2025-01-01")
doc1.add_paragraph("End Date: 2025-12-31")
doc1.save(co1_dir / "agreement.docx")

# 2. Invoice (.xlsx)
wb1 = openpyxl.Workbook()
ws1 = wb1.active
ws1.append(["Invoice Number", "Amount", "Company Name", "Date"])
ws1.append(["INV-2025-001", "500000", "Apollo Pharmacy", "2025-01-05"])
wb1.save(co1_dir / "invoice.xlsx")

# 3. Purchase Order (.docx)
po1 = Document()
po1.add_heading("PURCHASE ORDER", level=1)
po1.add_paragraph("PO Number: PO-998811")
po1.add_paragraph("Company Name: Apollo Pharmacy")
po1.add_paragraph("Total Value: INR 5,00,000")
po1.save(co1_dir / "purchase_order.docx")

# 4. Email (.docx)
em1 = Document()
em1.add_heading("EMAIL CORRESPONDENCE", level=1)
em1.add_paragraph("Sender: sales@msd.com")
em1.add_paragraph("Recipient: procurement@apollopharmacy.com")
em1.add_paragraph("Subject: Agreement & Invoice Confirmation for Apollo Pharmacy")
em1.save(co1_dir / "email.docx")


# --- COMPANY 2: Fortis Healthcare (INCOMPLETE / MISSING FIELDS FOR VALIDATION TEST) ---

# 1. Agreement (.docx) - missing start_date & end_date intentionally
doc2 = Document()
doc2.add_heading("AGREEMENT", level=1)
doc2.add_paragraph("Company Name: Fortis Healthcare")
doc2.add_paragraph("Agreement Value: INR 12,00,000")
doc2.save(co2_dir / "agreement.docx")

# 2. Invoice (.xlsx) - missing invoice_number
wb2 = openpyxl.Workbook()
ws2 = wb2.active
ws2.append(["Amount", "Company Name", "Date"])
ws2.append(["1200000", "Fortis Healthcare", "2025-01-10"])
wb2.save(co2_dir / "invoice.xlsx")

# 3. Purchase Order (.docx)
po2 = Document()
po2.add_heading("PURCHASE ORDER", level=1)
po2.add_paragraph("PO Number: PO-774422")
po2.add_paragraph("Company Name: Fortis Healthcare")
po2.add_paragraph("Total Value: INR 12,00,000")
po2.save(co2_dir / "purchase_order.docx")

# Note: email.docx is intentionally missing for Fortis Healthcare to test missing file detection!

# Create ZIP archive
with zipfile.ZipFile(zip_output, 'w', zipfile.ZIP_DEFLATED) as zipf:
    for root, dirs, files in os.walk(base_dir):
        for f in files:
            full_path = Path(root) / f
            arcname = full_path.relative_to(base_dir)
            zipf.write(full_path, arcname)

print(f"Sample validation ZIP generated successfully at: {zip_output}")
